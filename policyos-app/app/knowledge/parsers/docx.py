from io import BytesIO
from zipfile import BadZipFile, ZipFile

import docx
from docx import Document

from app.knowledge.schemas import ParsedDocument, ParsedSection, ParserError


class DocxDocumentParser:
    name = "python-docx"
    version = docx.__version__
    extensions = frozenset({".docx"})
    mime_types = frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            with ZipFile(BytesIO(content)) as archive:
                if any("/embeddings/" in name.lower() for name in archive.namelist()):
                    raise ParserError("DOCX embedded objects are not supported")
            document = Document(BytesIO(content))
        except ParserError:
            raise
        except (BadZipFile, KeyError, ValueError) as exc:
            raise ParserError("DOCX document could not be parsed") from exc

        sections: list[ParsedSection] = []
        current_title: str | None = None
        current: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.lower().startswith("heading"):
                if current or current_title:
                    sections.append(
                        ParsedSection(
                            index=len(sections) + 1,
                            title=current_title,
                            text="\n".join(current),
                        )
                    )
                current_title = text.strip() or None
                current = []
            elif text.strip():
                current.append(text)
        for table in document.tables:
            rows = [" | ".join(cell.text for cell in row.cells) for row in table.rows]
            if rows:
                current.append("\n".join(rows))
        if current or current_title:
            sections.append(
                ParsedSection(
                    index=len(sections) + 1,
                    title=current_title,
                    text="\n".join(current),
                )
            )
        if not sections or not any(section.text.strip() or section.title for section in sections):
            raise ParserError("DOCX document contains no extractable text")
        text = "\n\n".join(
            "\n".join(item for item in (section.title, section.text) if item)
            for section in sections
        )
        return ParsedDocument(
            text=text,
            sections=sections,
            parser_name=self.name,
            parser_version=self.version,
        )