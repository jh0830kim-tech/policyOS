from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document
from openpyxl import Workbook

from app.core.config import Settings
from app.knowledge.ingestion import normalize_filename, validate_document
from app.knowledge.normalization import TextNormalizer
from app.knowledge.parsers import (
    CsvDocumentParser,
    DocxDocumentParser,
    HwpDocumentParser,
    PdfDocumentParser,
    TextDocumentParser,
    XlsxDocumentParser,
)
from app.knowledge.schemas import (
    DocumentTooLargeError,
    IngestionRequest,
    InvalidDocumentError,
    ParserError,
    UnsupportedDocumentTypeError,
)


def request(filename: str, content: bytes, content_type: str = "text/plain") -> IngestionRequest:
    from uuid import uuid4

    return IngestionRequest(
        source_id=uuid4(), filename=filename, content_type=content_type, content=content
    )


def test_txt_and_markdown_ingestion_preserve_structured_text() -> None:
    parser = TextDocumentParser()
    txt = parser.parse(b"\xef\xbb\xbfFirst\r\nSecond", "brief.txt")
    markdown = parser.parse(b"# Heading\nBody\n## Next\nMore", "brief.md")
    assert txt.text == "First\r\nSecond"
    assert [section.title for section in markdown.sections] == ["Heading", "Next"]
    assert markdown.sections[0].text == "Body"


def test_pdf_preserves_page_numbers(monkeypatch) -> None:
    pages = [
        SimpleNamespace(extract_text=lambda: "Page one"),
        SimpleNamespace(extract_text=lambda: "Page two"),
    ]
    monkeypatch.setattr(
        "app.knowledge.parsers.pdf.PdfReader",
        lambda *_args, **_kwargs: SimpleNamespace(is_encrypted=False, pages=pages),
    )
    parsed = PdfDocumentParser().parse(b"%PDF-test", "law.pdf")
    assert parsed.page_count == 2
    assert [section.page_number for section in parsed.sections] == [1, 2]


def test_pdf_without_extractable_text_requires_explicit_ocr(monkeypatch) -> None:
    monkeypatch.setattr(
        "app.knowledge.parsers.pdf.PdfReader",
        lambda *_args, **_kwargs: SimpleNamespace(
            is_encrypted=False, pages=[SimpleNamespace(extract_text=lambda: "")]
        ),
    )
    with pytest.raises(ParserError, match="OCR is not automatic"):
        PdfDocumentParser().parse(b"%PDF-test", "scan.pdf")


def test_docx_extracts_headings_paragraphs_and_tables() -> None:
    document = Document()
    document.add_heading("Policy", level=1)
    document.add_paragraph("Paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    buffer = BytesIO()
    document.save(buffer)
    parsed = DocxDocumentParser().parse(buffer.getvalue(), "policy.docx")
    assert parsed.sections[0].title == "Policy"
    assert "Paragraph" in parsed.text and "A | B" in parsed.text


def test_csv_preserves_header_and_enforces_limits() -> None:
    parsed = CsvDocumentParser(max_rows=3, max_columns=2).parse(b"name,value\na,1", "data.csv")
    assert parsed.metadata["header"] == ["name", "value"]
    with pytest.raises(ParserError, match="row limit"):
        CsvDocumentParser(max_rows=1).parse(b"name,value\na,1", "data.csv")


def test_xlsx_preserves_sheet_metadata_and_formula_text() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Budget"
    sheet.append(["amount", "total"])
    sheet.append([10, "=SUM(A2:A2)"])
    buffer = BytesIO()
    workbook.save(buffer)
    parsed = XlsxDocumentParser().parse(buffer.getvalue(), "budget.xlsx")
    assert parsed.sheet_count == 1
    assert parsed.sections[0].sheet_name == "Budget"
    assert "=SUM(A2:A2)" in parsed.text


def test_hwp_adapter_fails_without_guessing_text() -> None:
    with pytest.raises(UnsupportedDocumentTypeError):
        HwpDocumentParser().parse(b"not parsed", "document.hwp")


def test_empty_oversized_executable_and_mismatched_files_are_rejected() -> None:
    settings = Settings(_env_file=None, knowledge_max_upload_bytes=4)
    with pytest.raises(InvalidDocumentError, match="Empty"):
        validate_document(request("empty.txt", b""), settings)
    with pytest.raises(DocumentTooLargeError):
        validate_document(request("large.txt", b"12345"), settings)
    with pytest.raises(InvalidDocumentError, match="Executable"):
        validate_document(request("payload.exe.txt", b"safe"), Settings(_env_file=None))
    with pytest.raises(InvalidDocumentError, match="does not match PDF"):
        validate_document(
            request("fake.pdf", b"plain", "application/pdf"), Settings(_env_file=None)
        )


def test_path_traversal_filename_is_reduced_to_safe_basename() -> None:
    assert normalize_filename("../../secret/../policy.txt") == "policy.txt"
    assert normalize_filename("..\\..\\budget.csv") == "budget.csv"

def test_normalization_is_unicode_safe_and_conservative() -> None:
    normalized = TextNormalizer().normalize("ＡＢＣ  \r\n\r\n\r\n  evidence   text")
    assert normalized == "ＡＢＣ\n\nevidence text"
